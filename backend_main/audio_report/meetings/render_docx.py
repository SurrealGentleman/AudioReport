from docx import Document
import copy


def render_docx_template(template_path, context, output_path):
    doc = Document(template_path)

    # Заменяет обычные переменные {{ key }} в тексте
    def replace_paragraph_variables(paragraph):
        for key, value in context.items():
            if isinstance(value, str):  # только строковые значения
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value)

    # Обработка обычных параграфов (без таблиц)
    for paragraph in doc.paragraphs:
        replace_paragraph_variables(paragraph)

    # Обработка таблиц
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            cell_texts = [cell.text.strip() for cell in row.cells]
            # ищем строку-шаблон для tasks
            if any('{{ tasks.' in cell_text for cell_text in cell_texts):
                task_template_row = row
                task_data = context.get("tasks", [])

                # Удаляем строку-шаблон из таблицы
                table._tbl.remove(task_template_row._tr)

                # Создаем строки по шаблону
                for task in task_data:
                    new_row = copy.deepcopy(task_template_row)
                    for cell in new_row.cells:
                        for key, value in task.items():
                            cell.text = cell.text.replace(f"{{{{ tasks.{key} }}}}", str(value))
                    table._tbl.append(new_row._tr)
                break  # только одна строка шаблона
            else:
                # обычная замена для других ячеек
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_paragraph_variables(paragraph)

    doc.save(output_path)
